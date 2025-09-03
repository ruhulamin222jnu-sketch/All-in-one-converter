from flask import Flask, render_template, request, send_file, jsonify
import os
from werkzeug.utils import secure_filename
from PIL import Image
from docx import Document
import pandas as pd
from pptx import Presentation
import shutil
import pdfplumber
import fitz  # PyMuPDF
from fpdf import FPDF
from pdfkit import from_file, configuration

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['DOWNLOAD_FOLDER'] = 'downloads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['DOWNLOAD_FOLDER'], exist_ok=True)

# wkhtmltopdf config (update your path)
pdfkit_config = configuration(wkhtmltopdf=r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe")


@app.route('/')
def index():
    return render_template('index.html')


def save_upload(file):
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)
    return filename, filepath


def cleanup_temp(files):
    for f in files:
        if os.path.exists(f):
            os.remove(f)


# ----------------- DOCX → PDF -----------------
@app.route('/word_to_pdf', methods=['POST'])
def word_to_pdf():
    file = request.files['file']
    filename, filepath = save_upload(file)

    pdf_path = os.path.join(app.config['DOWNLOAD_FOLDER'], filename.rsplit('.',1)[0]+'.pdf')

    # Convert DOCX → HTML → PDF
    doc = Document(filepath)
    html_content = "<html><body>"
    for para in doc.paragraphs:
        html_content += f"<p>{para.text}</p>"
    html_content += "</body></html>"

    html_file = os.path.join(app.config['UPLOAD_FOLDER'], filename.rsplit('.',1)[0]+".html")
    with open(html_file, "w", encoding="utf-8") as f:
        f.write(html_content)

    from_file(html_file, pdf_path, configuration=pdfkit_config)
    cleanup_temp([html_file])

    return send_file(pdf_path, as_attachment=True)


# ----------------- PDF → DOCX -----------------
@app.route('/pdf_to_word', methods=['POST'])
def pdf_to_word():
    file = request.files['file']
    filename, filepath = save_upload(file)
    doc_path = os.path.join(app.config['DOWNLOAD_FOLDER'], filename.rsplit('.',1)[0]+'.docx')
    doc = Document()

    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    doc.add_paragraph(text)
        doc.save(doc_path)
    except Exception as e:
        return jsonify({"error": str(e)}), 400

    return send_file(doc_path, as_attachment=True)


# ----------------- JPG → PDF -----------------
@app.route('/img_to_pdf', methods=['POST'])
def img_to_pdf():
    file = request.files['file']
    filename, filepath = save_upload(file)
    pdf_path = os.path.join(app.config['DOWNLOAD_FOLDER'], filename.rsplit('.',1)[0]+'.pdf')
    image = Image.open(filepath).convert('RGB')
    image.save(pdf_path, "PDF")
    return send_file(pdf_path, as_attachment=True)


# ----------------- PDF → JPG -----------------
@app.route('/pdf_to_img', methods=['POST'])
def pdf_to_img():
    file = request.files['file']
    filename, filepath = save_upload(file)
    pdf_doc = fitz.open(filepath)
    img_folder = os.path.join(app.config['DOWNLOAD_FOLDER'], filename.rsplit('.',1)[0])
    os.makedirs(img_folder, exist_ok=True)
    for i, page in enumerate(pdf_doc):
        pix = page.get_pixmap()
        pix.save(os.path.join(img_folder, f"page_{i+1}.png"))
    zip_path = os.path.join(app.config['DOWNLOAD_FOLDER'], filename.rsplit('.',1)[0]+'.zip')
    shutil.make_archive(zip_path.rsplit('.',1)[0], 'zip', img_folder)
    shutil.rmtree(img_folder)
    return send_file(zip_path, as_attachment=True)


# ----------------- PowerPoint → PDF -----------------
@app.route('/ppt_to_pdf', methods=['POST'])
def ppt_to_pdf():
    file = request.files['file']
    filename, filepath = save_upload(file)
    prs = Presentation(filepath)
    pdf_path = os.path.join(app.config['DOWNLOAD_FOLDER'], filename.rsplit('.',1)[0]+'.pdf')

    pdf = FPDF()
    for i, slide in enumerate(prs.slides):
        pdf.add_page()
        # Convert slide shapes to image (if needed, can extend with Pillow)
    pdf.output(pdf_path)
    return send_file(pdf_path, as_attachment=True)


# ----------------- Excel → PDF -----------------
@app.route('/excel_to_pdf', methods=['POST'])
def excel_to_pdf():
    file = request.files['file']
    filename, filepath = save_upload(file)
    df = pd.read_excel(filepath)
    html_file = os.path.join(app.config['UPLOAD_FOLDER'], filename.rsplit('.',1)[0]+'.html')
    df.to_html(html_file)
    pdf_path = os.path.join(app.config['DOWNLOAD_FOLDER'], filename.rsplit('.',1)[0]+'.pdf')
    from_file(html_file, pdf_path, configuration=pdfkit_config)
    cleanup_temp([html_file])
    return send_file(pdf_path, as_attachment=True)


# ----------------- CSV → DOCX -----------------
@app.route('/csv_to_doc', methods=['POST'])
def csv_to_doc():
    file = request.files['file']
    filename, filepath = save_upload(file)
    df = pd.read_csv(filepath)
    doc = Document()
    for _, row in df.iterrows():
        doc.add_paragraph(str(row.to_dict()))
    doc_path = os.path.join(app.config['DOWNLOAD_FOLDER'], filename.rsplit('.',1)[0]+'.docx')
    doc.save(doc_path)
    return send_file(doc_path, as_attachment=True)


# ----------------- DOCX → CSV -----------------
@app.route('/doc_to_csv', methods=['POST'])
def doc_to_csv():
    file = request.files['file']
    filename, filepath = save_upload(file)
    document = Document(filepath)
    data = [[para.text] for para in document.paragraphs]
    df = pd.DataFrame(data)
    csv_path = os.path.join(app.config['DOWNLOAD_FOLDER'], filename.rsplit('.',1)[0]+'.csv')
    df.to_csv(csv_path, index=False)
    return send_file(csv_path, as_attachment=True)


# ----------------- JPG → PNG -----------------
@app.route('/jpg_to_png', methods=['POST'])
def jpg_to_png():
    file = request.files['file']
    filename, filepath = save_upload(file)
    img_path = os.path.join(app.config['DOWNLOAD_FOLDER'], filename.rsplit('.',1)[0]+'.png')
    image = Image.open(filepath)
    image.save(img_path, "PNG")
    return send_file(img_path, as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True)
